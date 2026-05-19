import re
import json
import traceback
import numpy as np

# ── Load BERT model ONCE at module level ─────────────────────────
_bert_model = None

def get_bert_model():
    global _bert_model
    if _bert_model is None:
        try:
            print("[bert] Loading model for the first time...")
            from sentence_transformers import SentenceTransformer
            _bert_model = SentenceTransformer('all-MiniLM-L6-v2')
            print("[bert] Model loaded and cached!")
        except Exception as e:
            print(f"[bert] Failed to load model: {e}")
            _bert_model = None
    return _bert_model


# ══════════════════════════════════════════════════════════════════
# TEXT EXTRACTION
# ══════════════════════════════════════════════════════════════════

def extract_text(file_path):
    path = str(file_path).lower()
    try:
        if path.endswith('.pdf'):
            return _extract_pdf(file_path)
        elif path.endswith('.docx') or path.endswith('.doc'):
            return _extract_docx(file_path)
        else:
            with open(file_path, 'r', errors='ignore') as f:
                return f.read()
    except Exception as e:
        print(f"[extract_text] Error: {e}")
        return ""


def _extract_pdf(file_path):
    text = ""
    # Method 1 — PyPDF2
    try:
        import PyPDF2
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text += page.extract_text() or ""
        if len(text.strip()) > 100:
            print(f"[pdf] PyPDF2 extracted {len(text)} chars")
            return text
    except Exception as e:
        print(f"[pdf] PyPDF2 failed: {e}")
    # Method 2 — pdfminer
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(str(file_path))
        if text and len(text.strip()) > 100:
            print(f"[pdf] pdfminer extracted {len(text)} chars")
            return text
    except ImportError:
        pass
    except Exception as e:
        print(f"[pdf] pdfminer failed: {e}")
    print(f"[pdf] WARNING: Only {len(text)} chars extracted")
    return text


def _extract_docx(file_path):
    try:
        import docx
        doc = docx.Document(file_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        text = "\n".join(parts)
        print(f"[docx] Extracted {len(text)} chars")
        return text
    except Exception as e:
        print(f"[docx] Error: {e}")
        return ""


def preprocess_text(text):
    text = text.lower()
    text = re.sub(r'[^\w\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


# ══════════════════════════════════════════════════════════════════
# SKILL VOCABULARY
# ══════════════════════════════════════════════════════════════════

SKILL_VOCAB = [
    # Programming Languages
    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby',
    'swift', 'kotlin', 'go', 'rust', 'php', 'scala', 'r', 'matlab',
    'perl', 'haskell', 'elixir', 'dart', 'groovy',
    # Web Frontend
    'html', 'css', 'react', 'angular', 'vue', 'nextjs', 'nuxt',
    'bootstrap', 'tailwind', 'webpack', 'jquery', 'sass', 'redux',
    # Web Backend
    'nodejs', 'django', 'flask', 'fastapi', 'spring', 'express',
    'laravel', 'rails', 'asp.net', 'graphql', 'rest api', 'microservices',
    # AI / ML
    'machine learning', 'deep learning', 'nlp', 'computer vision',
    'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas',
    'numpy', 'matplotlib', 'seaborn', 'data analysis', 'data science',
    'neural network', 'bert', 'transformers', 'opencv', 'huggingface',
    'reinforcement learning', 'feature engineering', 'model deployment',
    'natural language processing', 'text mining', 'sentiment analysis',
    # Databases
    'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
    'sqlite', 'oracle', 'cassandra', 'dynamodb', 'firebase', 'neo4j',
    # Cloud & DevOps
    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'ci/cd', 'jenkins',
    'git', 'github', 'linux', 'bash', 'terraform', 'ansible', 'nginx',
    'devops', 'cloud computing', 'serverless',
    # Mobile
    'android', 'ios', 'flutter', 'react native', 'xamarin',
    # Data & Analytics
    'power bi', 'tableau', 'excel', 'hadoop', 'spark', 'kafka',
    'data visualization', 'statistics', 'big data', 'etl',
    # Soft Skills
    'communication', 'teamwork', 'leadership', 'problem solving',
    'project management', 'agile', 'scrum', 'time management',
    'critical thinking', 'collaboration', 'presentation',
]


def extract_skills(text):
    text_lower = text.lower()
    found = []
    for skill in SKILL_VOCAB:
        if skill in text_lower and skill not in found:
            found.append(skill)
    return found


# ══════════════════════════════════════════════════════════════════
# FACTOR 1 — REQUIRED SKILLS MATCH (35%)
# ══════════════════════════════════════════════════════════════════

def compute_required_skills_score(resume_text, required_skills_list):
    """
    Most important factor.
    Checks what % of MUST HAVE skills the candidate has.
    Missing even one required skill = big penalty.
    Returns 0.0 to 1.0
    """
    if not required_skills_list:
        # No required skills defined — fall back to general skill extraction
        return None

    resume_lower = resume_text.lower()
    matched = []
    missing = []

    for skill in required_skills_list:
        if skill.lower() in resume_lower:
            matched.append(skill)
        else:
            missing.append(skill)

    ratio = len(matched) / len(required_skills_list)
    print(f"[req_skills] {len(matched)}/{len(required_skills_list)} matched = {ratio:.4f}")
    print(f"[req_skills] Matched: {matched}")
    print(f"[req_skills] Missing: {missing}")
    return round(ratio, 4), matched, missing


# ══════════════════════════════════════════════════════════════════
# FACTOR 2 — OPTIONAL SKILLS BONUS (10%)
# ══════════════════════════════════════════════════════════════════

def compute_optional_skills_score(resume_text, optional_skills_list):
    """
    Bonus points for having nice-to-have skills.
    Returns 0.0 to 1.0
    """
    if not optional_skills_list:
        return 0.5  # neutral if no optional skills defined

    resume_lower = resume_text.lower()
    matched = [s for s in optional_skills_list if s.lower() in resume_lower]
    ratio = len(matched) / len(optional_skills_list)
    print(f"[opt_skills] {len(matched)}/{len(optional_skills_list)} matched = {ratio:.4f}")
    return round(ratio, 4), matched


# ══════════════════════════════════════════════════════════════════
# FACTOR 3 — BERT SEMANTIC SIMILARITY (25%)
# ══════════════════════════════════════════════════════════════════

def get_embedding(model, text, chunk_size=500, overlap=100):
    text = text.strip()
    if not text:
        return None
    words = text.split()
    if len(words) <= chunk_size:
        return model.encode(text, convert_to_tensor=False)
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = ' '.join(words[i:i + chunk_size])
        chunks.append(chunk)
        if i + chunk_size >= len(words):
            break
    embeddings = model.encode(chunks, convert_to_tensor=False)
    return np.mean(embeddings, axis=0)


def compute_bert_similarity(resume_text, job_text):
    try:
        model = get_bert_model()
        if model is None:
            return 0.0
        emb1 = get_embedding(model, resume_text)
        emb2 = get_embedding(model, job_text)
        if emb1 is None or emb2 is None:
            return 0.0
        dot   = np.dot(emb1, emb2)
        norm1 = np.linalg.norm(emb1)
        norm2 = np.linalg.norm(emb2)
        if norm1 == 0 or norm2 == 0:
            return 0.0
        score = float(dot / (norm1 * norm2))
        score = max(0.0, min(1.0, score))
        print(f"[bert] Score: {score:.4f}")
        return round(score, 4)
    except Exception as e:
        print(f"[bert] ERROR: {e}")
        traceback.print_exc()
        return 0.0


# ══════════════════════════════════════════════════════════════════
# FACTOR 4 — EXPERIENCE MATCH (15%)
# ══════════════════════════════════════════════════════════════════

def extract_experience_years(text):
    text_lower = text.lower()

    fresher_keywords = ['fresher', 'fresh graduate', 'no experience',
                        'recently graduated', '0 years', 'zero experience']
    is_fresher = any(kw in text_lower for kw in fresher_keywords)
    has_internship = any(kw in text_lower for kw in [
        'intern', 'internship', 'trainee', 'training', 'part time', 'freelance'])

    # Range pattern: "0-2 years"
    m = re.search(r'(\d+)\s*(?:to|-)\s*(\d+)\s*\+?\s*years?\s*(?:of\s+)?(?:experience|exp)',
                  text_lower)
    if m:
        years = (float(m.group(1)) + float(m.group(2))) / 2
        print(f"[exp] Range {m.group(1)}-{m.group(2)} → {years} years")
        return years

    # Exact patterns
    patterns = [
        r'(\d+\.?\d*)\s*\+?\s*years?\s+(?:of\s+)?(?:experience|work|industry)',
        r'(\d+\.?\d*)\s*\+?\s*years?\s+(?:professional|relevant|total)',
        r'experience\s+of\s+(\d+\.?\d*)\s*\+?\s*years?',
        r'(\d+\.?\d*)\s*\+?\s*yrs?\s+(?:of\s+)?exp',
        r'over\s+(\d+\.?\d*)\s*years?',
        r'more\s+than\s+(\d+\.?\d*)\s*years?',
    ]
    for pat in patterns:
        m = re.search(pat, text_lower)
        if m:
            years = float(m.group(1))
            print(f"[exp] Found {years} years")
            return years

    # Months
    m = re.search(r'(\d+)\s*months?\s+(?:of\s+)?(?:experience|internship|work)', text_lower)
    if m:
        years = round(float(m.group(1)) / 12, 1)
        print(f"[exp] {m.group(1)} months → {years} years")
        return years

    if is_fresher and has_internship:
        return 0.5
    if is_fresher:
        return 0.0

    return None


def compute_experience_score(resume_text, min_experience):
    """
    Compare candidate's experience against job requirement.
    Returns 0.0 to 1.0
    """
    exp_years = extract_experience_years(resume_text)

    # Map min_experience field to required years
    req_map = {
        'any': 0, 'fresher': 0, '1': 1, '2': 2,
        '3': 3, '5': 5, '8': 8
    }
    required_years = req_map.get(str(min_experience), 0)

    job_wants_fresher = (min_experience == 'fresher')

    if exp_years is None:
        # Experience not mentioned
        if required_years == 0:
            return 0.70  # neutral for no-requirement jobs
        return 0.55  # slight uncertainty

    if job_wants_fresher:
        if exp_years <= 1:
            score = 1.0   # perfect — fresher for fresher job
        elif exp_years <= 3:
            score = 0.70  # slightly overqualified
        else:
            score = 0.50  # overqualified
    elif required_years == 0:
        # Any experience acceptable
        score = min(1.0, 0.60 + (exp_years * 0.08))
    else:
        if exp_years >= required_years:
            # Meets or exceeds requirement
            excess = exp_years - required_years
            if excess > 5:
                score = 0.85  # overqualified — slight penalty
            else:
                score = 1.0
        else:
            # Below requirement — proportional penalty
            ratio = exp_years / required_years
            score = max(0.20, ratio * 0.85)

    print(f"[exp] candidate={exp_years}yr required={required_years}yr score={score:.4f}")
    return round(score, 4)


# ══════════════════════════════════════════════════════════════════
# FACTOR 5 — EDUCATION + CGPA (15%)
# ══════════════════════════════════════════════════════════════════

DEGREE_LEVEL = {
    'phd': 7, 'mtech': 6, 'mca': 5, 'mba': 5, 'msc': 5,
    'btech': 4, 'bca': 3, 'bsc': 3, 'bcom': 3, 'ba': 3,
    'diploma': 2, '12th': 1, '10th': 0,
}

DEGREE_KEYWORDS = {
    'phd':     ['phd', 'ph.d', 'doctorate'],
    'mtech':   ['m.tech', 'mtech', 'master of technology', 'm.e.'],
    'mca':     ['mca', 'master of computer'],
    'mba':     ['mba', 'master of business'],
    'msc':     ['m.sc', 'msc', 'master of science'],
    'btech':   ['b.tech', 'btech', 'bachelor of technology', 'b.e.', ' be '],
    'bca':     ['bca', 'bachelor of computer'],
    'bsc':     ['b.sc', 'bsc', 'bachelor of science'],
    'bcom':    ['b.com', 'bcom', 'bachelor of commerce'],
    'ba':      ['b.a.', ' ba ', 'bachelor of arts'],
    'diploma': ['diploma', 'polytechnic'],
    '12th':    ['12th', 'intermediate', 'higher secondary', 'hsc'],
    '10th':    ['10th', 'matriculation', 'ssc'],
}

MIN_EDU_LEVEL = {
    'any': 0, '10th': 0, '12th': 1, 'diploma': 2,
    'bca': 3, 'btech': 4, 'mca': 5, 'mtech': 6,
    'mba': 5, 'phd': 7,
}


def extract_candidate_degree(text):
    text_lower = text.lower()
    best_level = -1
    best_deg   = None
    for deg, keywords in DEGREE_KEYWORDS.items():
        if any(kw in text_lower for kw in keywords):
            level = DEGREE_LEVEL.get(deg, 0)
            if level > best_level:
                best_level = level
                best_deg   = deg
    return best_deg, best_level


def extract_cgpa(text):
    text_lower = text.lower()

    # CGPA x/10
    m = re.search(r'cgpa[\s:]*(\d+\.?\d*)\s*/\s*10', text_lower)
    if m:
        return float(m.group(1)) / 10.0

    # GPA x/4
    m = re.search(r'gpa[\s:]*(\d+\.?\d*)\s*/\s*4', text_lower)
    if m:
        return float(m.group(1)) / 4.0

    # x.x/10
    m = re.search(r'(\d+\.?\d*)\s*/\s*10', text_lower)
    if m:
        val = float(m.group(1))
        if 0 < val <= 10:
            return val / 10.0

    # Percentage
    m = re.search(r'(\d{2,3}\.?\d*)\s*%', text_lower)
    if m:
        val = float(m.group(1))
        if 0 < val <= 100:
            return val / 100.0

    # "scored 85" or "aggregate 78"
    m = re.search(r'(?:scored|aggregate|marks|percentage)[\s:]*(\d{2,3}\.?\d*)', text_lower)
    if m:
        val = float(m.group(1))
        if 0 < val <= 100:
            return val / 100.0

    return None


def compute_education_score(resume_text, min_education, min_cgpa):
    """
    Combined education + CGPA score.
    Returns 0.0 to 1.0
    """
    candidate_deg, candidate_level = extract_candidate_degree(resume_text)
    required_level = MIN_EDU_LEVEL.get(min_education, 0)
    cgpa_normalized = extract_cgpa(resume_text)

    # ── Degree score ──────────────────────────────────────────
    if required_level == 0:
        degree_score = 1.0  # any degree accepted
    elif candidate_level >= required_level:
        degree_score = 1.0  # meets or exceeds
    elif candidate_level == required_level - 1:
        degree_score = 0.65  # one level below
    elif candidate_level >= 0:
        degree_score = 0.35  # significantly below
    else:
        degree_score = 0.50  # unknown — neutral

    # ── CGPA score ────────────────────────────────────────────
    if cgpa_normalized is None:
        cgpa_score = 0.65  # not found — neutral
    elif min_cgpa and min_cgpa > 0:
        cutoff = min_cgpa / 10.0
        if cgpa_normalized >= cutoff:
            # Above cutoff — reward proportionally
            cgpa_score = min(1.0, 0.70 + (cgpa_normalized - cutoff) * 2)
        else:
            # Below cutoff — penalty
            cgpa_score = max(0.0, (cgpa_normalized / cutoff) * 0.50)
    else:
        # No cutoff — reward good CGPA
        if cgpa_normalized >= 0.85:
            cgpa_score = 1.0
        elif cgpa_normalized >= 0.75:
            cgpa_score = 0.85
        elif cgpa_normalized >= 0.60:
            cgpa_score = 0.70
        else:
            cgpa_score = 0.50

    # Combined: 60% degree + 40% CGPA
    combined = round(0.60 * degree_score + 0.40 * cgpa_score, 4)
    print(f"[edu] degree={candidate_deg}(lvl {candidate_level}) req_lvl={required_level} "
          f"deg_score={degree_score:.2f} cgpa={cgpa_normalized} cgpa_score={cgpa_score:.2f} "
          f"combined={combined:.4f}")
    return combined


# ══════════════════════════════════════════════════════════════════
# FACTOR 6 — RESUME COMPLETENESS (5%)
# ══════════════════════════════════════════════════════════════════

def compute_completeness_score(resume_text):
    """
    Check if resume has all standard sections.
    Returns 0.0 to 1.0
    """
    text_lower = resume_text.lower()
    score = 0
    total = 6

    # Contact info
    if re.search(r'[\w.]+@[\w.]+\.\w+', resume_text) or re.search(r'\d{10}', resume_text):
        score += 1

    # Education section
    if any(kw in text_lower for kw in ['education', 'qualification', 'academic', 'degree', 'university', 'college']):
        score += 1

    # Experience / Projects
    if any(kw in text_lower for kw in ['experience', 'work', 'employment', 'project', 'internship']):
        score += 1

    # Skills section
    if any(kw in text_lower for kw in ['skill', 'technical', 'technology', 'tools', 'proficient']):
        score += 1

    # Summary / Objective
    if any(kw in text_lower for kw in ['summary', 'objective', 'profile', 'about', 'overview']):
        score += 1

    # Certifications / Achievements
    if any(kw in text_lower for kw in ['certif', 'achievement', 'award', 'accomplishment', 'course']):
        score += 1

    result = round(score / total, 4)
    print(f"[completeness] {score}/{total} sections found = {result:.4f}")
    return result


# ══════════════════════════════════════════════════════════════════
# AUTO REJECTION LOGIC
# ══════════════════════════════════════════════════════════════════

def check_auto_rejection(resume_text, job, req_matched, req_total, cgpa_normalized, candidate_level):
    """
    Check if candidate should be auto rejected based on hard filters.
    Returns (should_reject: bool, reason: str)
    """
    reasons = []

    # Rule 1 — Zero required skills matched
    if req_total > 0 and req_matched == 0:
        reasons.append(f"0 out of {req_total} required skills found in resume")

    # Rule 2 — CGPA below minimum cutoff
    if job.min_cgpa and job.min_cgpa > 0 and cgpa_normalized is not None:
        cutoff = job.min_cgpa / 10.0
        if cgpa_normalized < cutoff:
            reasons.append(
                f"CGPA {cgpa_normalized*10:.1f} is below minimum cutoff of {job.min_cgpa}"
            )

    # Rule 3 — Education well below requirement
    required_level = MIN_EDU_LEVEL.get(job.min_education, 0)
    if required_level > 0 and candidate_level >= 0:
        if candidate_level < required_level - 1:
            reasons.append(
                f"Education level ({candidate_level}) is below minimum requirement"
            )

    # Rule 4 — Required skills match below 20% for jobs with 5+ required skills
    if req_total >= 5 and req_matched / req_total < 0.20:
        reasons.append(
            f"Only {req_matched}/{req_total} required skills matched (below 20% threshold)"
        )

    if reasons:
        return True, " | ".join(reasons)
    return False, ""


# ══════════════════════════════════════════════════════════════════
# MAIN ANALYSIS FUNCTION
# ══════════════════════════════════════════════════════════════════

def analyze_resume(file_path, job_description, job=None):
    """
    Full 6-factor resume analysis.

    Scoring Formula:
        Required Skills Match   35%
        BERT Semantic           25%
        Experience Match        15%
        Education + CGPA        15%
        Optional Skills Bonus   10%
        Resume Completeness      5%  (included in education factor)
    """
    print(f"\n{'='*50}")
    print(f"[analyze] Processing: {file_path}")

    # ── Extract text ──────────────────────────────────────────
    raw_text = extract_text(file_path)
    if not raw_text.strip():
        print("[analyze] WARNING: No text extracted!")
        raw_text = "No text extracted"
    print(f"[analyze] Extracted {len(raw_text)} characters")

    # ── Get structured job fields ─────────────────────────────
    if job:
        required_skills_list = job.required_skills_list
        optional_skills_list = job.optional_skills_list
        min_experience       = job.min_experience
        min_education        = job.min_education
        min_cgpa             = job.min_cgpa
    else:
        required_skills_list = []
        optional_skills_list = []
        min_experience       = 'any'
        min_education        = 'any'
        min_cgpa             = 0.0

    # ── Compute all factors ───────────────────────────────────

    # Factor 1 — Required Skills (35%)
    req_result = compute_required_skills_score(raw_text, required_skills_list)
    if req_result is not None:
        req_score, req_matched_list, req_missing_list = req_result
        req_matched_count = len(req_matched_list)
        req_total_count   = len(required_skills_list)
    else:
        # No required skills defined — use general skill extraction
        general_skills = extract_skills(raw_text)
        job_skills     = extract_skills(job_description)
        if job_skills:
            matched = [s for s in job_skills if s in general_skills]
            req_score = len(matched) / len(job_skills)
        else:
            req_score = 0.60
        req_matched_count = 0
        req_total_count   = 0
        req_missing_list  = []

    # Factor 2 — Optional Skills (10%)
    opt_result = compute_optional_skills_score(raw_text, optional_skills_list)
    if isinstance(opt_result, tuple):
        opt_score, opt_matched_list = opt_result
        opt_matched_count = len(opt_matched_list)
    else:
        opt_score = opt_result
        opt_matched_count = 0

    # Factor 3 — BERT Semantic (25%)
    bert_score = compute_bert_similarity(raw_text, job_description)

    # Factor 4 — Experience (15%)
    exp_score = compute_experience_score(raw_text, min_experience)

    # Factor 5 — Education + CGPA (15%)
    edu_score = compute_education_score(raw_text, min_education, min_cgpa)

    # Factor 6 — Resume Completeness (5% — included in final)
    completeness = compute_completeness_score(raw_text)

    # ── Extract CGPA and degree for auto-rejection check ──────
    cgpa_normalized  = extract_cgpa(raw_text)
    _, candidate_lvl = extract_candidate_degree(raw_text)

    # ── Auto rejection check ──────────────────────────────────
    auto_reject  = False
    reject_reason = ""
    if job:
        auto_reject, reject_reason = check_auto_rejection(
            raw_text, job,
            req_matched_count, req_total_count,
            cgpa_normalized, candidate_lvl
        )

    # ── Final weighted score ──────────────────────────────────
    final_score = round(
        0.35 * req_score    +
        0.25 * bert_score   +
        0.15 * exp_score    +
        0.15 * edu_score    +
        0.10 * opt_score,
        4
    )

    # Add completeness as small bonus (up to +3%)
    final_score = min(1.0, round(final_score + (completeness * 0.03), 4))

    # Auto rejected candidates get capped at 30%
    if auto_reject:
        final_score = min(final_score, 0.30)

    # Extract all skills for display
    all_skills = extract_skills(raw_text)
    if required_skills_list:
        # Also include matched required/optional skills
        for s in req_matched_list if req_result else []:
            if s not in all_skills:
                all_skills.append(s)

    print(f"\n[analyze] SCORE BREAKDOWN:")
    print(f"  Required Skills:  {req_score*100:.1f}% (weight 35%)")
    print(f"  BERT Semantic:    {bert_score*100:.1f}% (weight 25%)")
    print(f"  Experience:       {exp_score*100:.1f}% (weight 15%)")
    print(f"  Education+CGPA:   {edu_score*100:.1f}% (weight 15%)")
    print(f"  Optional Skills:  {opt_score*100:.1f}% (weight 10%)")
    print(f"  Completeness:     {completeness*100:.1f}% (bonus +3%)")
    print(f"  AUTO REJECT:      {auto_reject} — {reject_reason}")
    print(f"  FINAL SCORE:      {final_score*100:.1f}%")
    print(f"{'='*50}\n")

    return {
        'raw_text':               raw_text,
        'tfidf_score':            req_score,      # stored as tfidf for display
        'bert_score':             bert_score,
        'final_score':            final_score,
        'skills':                 all_skills,
        'required_skills_score':  req_score,
        'optional_skills_score':  opt_score,
        'experience_score':       exp_score,
        'education_score':        edu_score,
        'completeness_score':     completeness,
        'required_skills_matched': req_matched_count,
        'optional_skills_matched': opt_matched_count,
        'auto_rejected':          auto_reject,
        'rejection_reason':       reject_reason,
    }


def rank_candidates(job):
    candidates = list(job.candidates.exclude(final_score=0).order_by('-final_score'))
    for i, c in enumerate(candidates, start=1):
        c.rank = i
        c.save(update_fields=['rank'])
    return candidates
